import re, sys, json
from pathlib import Path

NARANJA = "D57A23"
LEYENDA_IA = ("Este documento ha sido generado con apoyo de inteligencia artificial a partir de la "
              "información proporcionada. Debe ser revisado y validado por un profesional tributario "
              "antes de su uso.")
_ASSETS = Path(__file__).parent.parent / "assets"

class TokenInexistente(Exception): ...
class MontoLiteral(Exception): ...

def fmt_clp(n):
    return "$" + f"{int(n):,}".replace(",", ".")

def resolver_token(datos, token):
    # Los tokens viven en un bloque plano datos["tokens"]; no se navega f22[]/balance.
    entrada = datos.get("tokens", {}).get(token)
    if entrada is None:
        raise TokenInexistente(token)
    valor = entrada.get("valor") if isinstance(entrada, dict) else entrada
    if not isinstance(valor, int):
        raise TokenInexistente(f"{token} no resuelve a monto entero")
    return valor

# monto CLP: >=4 dígitos con separador de miles, o precedido de $.
_MONTO = re.compile(r"(?<!\{)\$?\d{1,3}(?:\.\d{3})+(?!\})|(?<!\{)\$\d{4,}(?!\})")

def _sin_placeholders(prosa):
    return re.sub(r"\{\{[^}]+\}\}", "", prosa)

def validar_sin_montos_literales(prosa):
    limpia = _sin_placeholders(prosa)
    m = _MONTO.search(limpia)
    if m:
        raise MontoLiteral(f"Monto literal en prosa (debió ser placeholder): '{m.group(0)}'")

def sustituir_placeholders(prosa, datos):
    validar_sin_montos_literales(prosa)
    def repl(m):
        return fmt_clp(resolver_token(datos, m.group(1).strip()))
    return re.sub(r"\{\{([^}]+)\}\}", repl, prosa)

def _campo_pagina(paragraph):
    """Inserta un campo PAGE (número de página dinámico) en el párrafo."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    ini = OxmlElement("w:fldChar"); ini.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
    run._r.append(ini); run._r.append(instr); run._r.append(fin)

def construir_docx(datos, secciones, salida):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Arial"; normal.size = Pt(11)
    # Encabezado con imagotipo
    header = doc.sections[0].header
    logo = _ASSETS / "imagotipo-principal-negro.png"
    if not logo.exists():
        raise FileNotFoundError(f"asset de marca faltante: {logo}")
    header.paragraphs[0].add_run().add_picture(str(logo), width=Inches(2.2))
    # Pie: leyenda confidencial + número de página + isologo naranja
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = "Documento de uso interno — Confidencial   |   Página "
    _campo_pagina(fp)
    iso = _ASSETS / "isologo-naranjo.png"
    if not iso.exists():
        raise FileNotFoundError(f"asset de marca faltante: {iso}")
    fp.add_run("   ").add_picture(str(iso), width=Inches(0.28))
    for sec in secciones:
        h = doc.add_heading(level=1)
        run = h.add_run((sec.get("titulo") or "").upper())
        run.font.name = "Arial"; run.font.color.rgb = RGBColor.from_string(NARANJA)
        prosa = sustituir_placeholders(sec.get("prosa", ""), datos)
        if prosa: doc.add_paragraph(prosa)
        tabla = sec.get("tabla")
        if tabla:
            t = doc.add_table(rows=len(tabla), cols=len(tabla[0]))
            try:
                t.style = "Light Grid Accent 1"
            except KeyError:
                pass
            for i, fila in enumerate(tabla):
                for j, val in enumerate(fila):
                    t.rows[i].cells[j].text = sustituir_placeholders(str(val), datos)
    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run(LEYENDA_IA); r.italic = True; r.font.size = Pt(8)
    doc.save(salida)
    return salida

def _fuente(dato):
    if isinstance(dato, dict) and "archivo" in dato:
        return f"{dato.get('archivo')} p.{dato.get('pagina')}: {dato.get('texto_cercano','')}"
    return ""

def construir_anexo(datos, salida):
    from openpyxl import Workbook
    wb = Workbook(); wb.remove(wb.active)
    # Ficha
    ws = wb.create_sheet("Ficha")
    ficha = datos.get("cliente", {})
    ws.append(["Campo", "Valor"])
    for k in ("razon_social", "rut", "regimen", "fecha_constitucion", "inicio_actividades"):
        ws.append([k, ficha.get(k)])
    # F22 / F29 con columna Fuente
    for hoja, key in (("F22", "f22"), ("F29", "f29")):
        ws = wb.create_sheet(hoja)
        ws.append(["Periodo/Año", "Código", "Valor", "Fuente"])
        for item in datos.get(key, []):
            etiqueta = item.get("anio") or item.get("periodo")
            for cod, dato in (item.get("codigos") or {}).items():
                if dato is None: continue
                ws.append([etiqueta, cod, dato.get("valor"), _fuente(dato)])
    # Balance — Resumen (solo escalares) + Cuentas (detalle), en hojas separadas.
    bal = datos.get("balance") or {}
    ws = wb.create_sheet("Balance Resumen")
    ws.append(["Campo", "Valor"])
    for k, v in bal.items():
        if str(k).startswith("_") or k == "cuentas": continue
        if isinstance(v, (list, dict)): continue   # nunca escribir estructuras crudas en una celda
        ws.append([k, v])
    ws = wb.create_sheet("Balance Cuentas")
    ws.append(["Código", "Descripción", "Activo", "Pasivo"])
    for c in (bal.get("cuentas") or []):
        ws.append([c.get("codigo"), c.get("descripcion"), c.get("activo"), c.get("pasivo")])
    # Reconciliaciones
    ws = wb.create_sheet("Reconciliaciones")
    ws.append(["Regla", "Esperado", "Obtenido", "Cuadra"])
    for r in datos.get("reconciliaciones", []):
        ws.append([r.get("regla"), r.get("esperado"), r.get("obtenido"), r.get("cuadra")])
    # Alertas
    ws = wb.create_sheet("Alertas")
    ws.append(["Nivel", "Código", "Periodo", "Detalle"])
    for a in datos.get("alertas_deterministas", []):
        ws.append([a.get("nivel"), a.get("codigo"), a.get("periodo"), a.get("detalle")])
    wb.save(salida)
    return salida

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Genera informe .docx + anexo .xlsx desde datos.json")
    ap.add_argument("datos", help="ruta a datos.json (incluye bloque 'tokens')")
    ap.add_argument("secciones", help="ruta a secciones.json (lista de {titulo, prosa, tabla})")
    ap.add_argument("--docx", default="informe.docx")
    ap.add_argument("--xlsx", default="anexo.xlsx")
    a = ap.parse_args()
    datos = json.load(open(a.datos, encoding="utf-8"))
    secciones = json.load(open(a.secciones, encoding="utf-8"))
    construir_docx(datos, secciones, a.docx)     # aborta ruidoso si hay monto literal o token faltante
    construir_anexo(datos, a.xlsx)
    print(json.dumps({"docx": a.docx, "xlsx": a.xlsx}, ensure_ascii=False))
