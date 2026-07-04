import importlib.util, os, pytest
from docx import Document
from openpyxl import load_workbook
def _load(name):
    spec = importlib.util.spec_from_file_location(name,
        os.path.join(os.path.dirname(__file__), "..", "scripts", name + ".py"))
    m = importlib.util.module_from_spec(spec); spec.loader.exec_module(m); return m
gi = _load("generar_informe")

DATOS = {"tokens": {"f22.2026.843": {"valor": 3186616079,
                                     "archivo": "F22.pdf", "pagina": 1, "texto_cercano": "843 ..."}}}

def test_fmt_clp():
    assert gi.fmt_clp(1234567) == "$1.234.567"

def test_sustituye_token():
    out = gi.sustituir_placeholders("Patrimonio: {{f22.2026.843}}.", DATOS)
    assert out == "Patrimonio: $3.186.616.079."

def test_token_inexistente_falla():
    with pytest.raises(gi.TokenInexistente):
        gi.sustituir_placeholders("X: {{f22.2026.999}}.", DATOS)

def test_monto_literal_en_prosa_falla():
    with pytest.raises(gi.MontoLiteral):
        gi.sustituir_placeholders("El total es $1.234.567 pesos.", DATOS)

def test_anio_y_porcentaje_no_falsos_positivos():
    # 2026 (año) y 27% (tasa) no deben gatillar la guardia
    out = gi.sustituir_placeholders("En 2026 la tasa IDPC es 27%.", DATOS)
    assert out == "En 2026 la tasa IDPC es 27%."

def test_construir_docx_marca_y_tabla(tmp_path):
    datos = {"tokens": {"f22.2026.843": {"valor": 900}}}
    secciones = [{"titulo": "Resumen", "prosa": "Patrimonio {{f22.2026.843}}.",
                  "tabla": [["Concepto", "Valor"], ["Patrimonio", "{{f22.2026.843}}"]]}]
    out = gi.construir_docx(datos, secciones, str(tmp_path / "informe.docx"))
    doc = Document(out)
    textos = "\n".join(p.text for p in doc.paragraphs)
    assert "RESUMEN" in textos.upper()
    assert "$900" in textos                      # placeholder sustituido en prosa
    celdas = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    assert "$900" in celdas                      # placeholder sustituido en celda de tabla
    assert any("profesional tributario" in p.text for p in doc.paragraphs)  # leyenda IA
    assert len(doc.tables) == 1
    footer = doc.sections[0].footer
    assert any("Confidencial" in p.text for p in footer.paragraphs)          # leyenda pie
    assert any("image" in r.reltype for r in footer.part.rels.values())      # isologo en pie
    header = doc.sections[0].header
    assert any("image" in r.reltype for r in header.part.rels.values())      # imagotipo en encabezado

def test_construir_docx_celda_monto_literal_aborta(tmp_path):
    # Un monto CLP tecleado en una celda (no placeholder) debe abortar la generación.
    datos = {"tokens": {"f22.2026.843": {"valor": 900}}}
    secciones = [{"titulo": "Resumen", "prosa": "",
                  "tabla": [["Concepto", "Valor"], ["Patrimonio", "$1.234.567"]]}]
    with pytest.raises(gi.MontoLiteral):
        gi.construir_docx(datos, secciones, str(tmp_path / "informe.docx"))

def test_anexo_hojas_y_fuente(tmp_path):
    datos = {"f22": [{"anio": 2026, "codigos": {
        "843": {"valor": 900, "archivo": "F22.pdf", "pagina": 1, "texto_cercano": "843 900 Patrimonio"}}}],
        "balance": {"total_activo_sii": 900, "cuentas": [
            {"codigo": "1.1", "descripcion": "CAJA", "activo": 900, "pasivo": 0}]}}
    out = gi.construir_anexo(datos, str(tmp_path / "anexo.xlsx"))
    wb = load_workbook(out)
    assert "F22" in wb.sheetnames
    assert "Balance Resumen" in wb.sheetnames and "Balance Cuentas" in wb.sheetnames
    ws = wb["F22"]
    filas = list(ws.iter_rows(values_only=True))
    assert filas[0][-1] == "Fuente"                       # última columna
    assert any("F22.pdf" in str(c) for row in filas for c in row if c)
    # ninguna celda contiene una lista o dict crudo
    for name in wb.sheetnames:
        for row in wb[name].iter_rows(values_only=True):
            for c in row:
                assert not isinstance(c, (list, dict))
